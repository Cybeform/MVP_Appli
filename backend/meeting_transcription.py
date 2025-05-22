# backend/meeting.py

from dotenv import load_dotenv
load_dotenv()  # charge OPENAI_API_KEY et HUGGINGFACE_TOKEN depuis .env

import os
import sys
import tempfile
import queue
import numpy as np
import openai
from concurrent.futures import ThreadPoolExecutor, as_completed

from pydub import AudioSegment
from docx import Document
from pyannote.audio import Pipeline

# ——————————————————————————————
# Configuration générale
# ——————————————————————————————
openai.api_key = os.getenv("OPENAI_API_KEY")
HF_TOKEN       = os.getenv("HUGGINGFACE_HUB_TOKEN") or os.getenv("HUGGINGFACE_TOKEN")
_MAX_BYTES     = 25 * 1024 * 1024      # 25 MiB max for Whisper upload
_CHUNK_MS       = 4 * 60 * 1000        # 4 minutes per chunk
_DIAR_THRESHOLD = 1        # 5 minutes max for diarization

# ——————————————————————————————
# Chargement du pipeline Pyannote pour diarization
# ——————————————————————————————
try:
    pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization",
        use_auth_token=HF_TOKEN
    )
    print("[diarization] ✅ Pipeline chargé avec succès.", file=sys.stderr)
except Exception as e:
    print(f"[diarization] ⚠️ Échec du chargement du pipeline : {e}", file=sys.stderr)
    pipeline = None

# ——————————————————————————————
# Variables & helpers pour enregistrement live
# ——————————————————————————————
_rec_queue   = queue.Queue()
_stream      = None
_output_file = None
_fs          = None

def start_recording(output_file: str = "recordings/meeting.wav",
                    fs: int = 44100,
                    channels: int = 1) -> None:
    """
    Démarre un enregistrement audio illimité, stocke dans `output_file`.
    """
    import sounddevice as sd
    from scipy.io.wavfile import write as wav_write

    global _stream, _rec_queue, _output_file, _fs
    _output_file = output_file
    _fs          = fs
    _rec_queue   = queue.Queue()
    os.makedirs(os.path.dirname(output_file), exist_ok=True)

    def callback(indata, frames, _, status):
        if status:
            print(f"[recording] {status}", file=sys.stderr)
        _rec_queue.put(indata.copy())

    _stream = sd.InputStream(
        samplerate=fs,
        channels=channels,
        callback=callback
    )
    _stream.start()
    print(f"[recording] Démarrage enregistrement → {_output_file}", file=sys.stderr)

def stop_recording() -> str:
    """
    Arrête l’enregistrement en cours et retourne le chemin WAV sauvegardé.
    """
    import sounddevice as sd
    from scipy.io.wavfile import write as wav_write

    global _stream, _rec_queue, _output_file, _fs
    if _stream is None:
        raise RuntimeError("Aucun enregistrement en cours.")
    _stream.stop()
    _stream.close()

    frames = []
    while not _rec_queue.empty():
        frames.append(_rec_queue.get())
    audio_data = np.concatenate(frames, axis=0)
    wav_write(_output_file, _fs, audio_data)
    print(f"[recording] Enregistrement sauvegardé → {_output_file}", file=sys.stderr)
    _stream = None
    return _output_file

# ——————————————————————————————
# Helpers pour transcription Whisper
# ——————————————————————————————
def _encode_and_transcribe(seg: AudioSegment) -> str:
    """
    Encode un segment Pydub en MP3 et l’envoie à Whisper.
    """
    with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as tmp:
        seg.export(tmp.name, format="mp3")
        path = tmp.name
    with open(path, "rb") as f:
        resp = openai.audio.transcriptions.create(file=f, model="whisper-1")
    os.remove(path)
    return resp.text

def _transcribe_simple(audio_input) -> str:
    """
    Transcrit un fichier audio (path ou AudioSegment), en découpant si >25 MiB.
    """
    if isinstance(audio_input, AudioSegment):
        audio = audio_input
        total_size = None
    else:
        audio = AudioSegment.from_file(audio_input)
        total_size = os.path.getsize(audio_input)

    if total_size is not None and total_size <= _MAX_BYTES:
        return _encode_and_transcribe(audio)

    texts = []
    duration_ms = len(audio)
    for start_ms in range(0, duration_ms, _CHUNK_MS):
        chunk = audio[start_ms:start_ms + _CHUNK_MS]
        texts.append(_encode_and_transcribe(chunk))
    return "\n".join(texts)

# ——————————————————————————————
# Générateur de progression + diarization limitée
# ——————————————————————————————
def transcribe_with_progress(audio_file: str):
    """
    Générateur d’événements SSE :
      - phase=diarization status=start|skipped|end count
      - phase=transcription total/done
      - phase=summary status=start|end
      - phase=docx status=start|end path
      - return (transcript, summary, docx_path)
    """
    # 1) ré-encode en WAV PCM pour Pyannote
    tmp_wav = tempfile.NamedTemporaryFile(suffix=".wav", delete=False).name
    AudioSegment.from_file(audio_file).export(tmp_wav, format="wav")

    # 2) diarization si <5min et pipeline dispo
    yield {"phase":"diarization","status":"start"}
    audio_seg = AudioSegment.from_file(tmp_wav)
    duration_ms = len(audio_seg)

    if pipeline is None or duration_ms > _DIAR_THRESHOLD:
        yield {"phase":"diarization","status":"skipped","count":1}
        segments = [(0, None)]
    else:
        diar = pipeline(tmp_wav)
        segments = [
            (int(turn.start * 1000), int(turn.end * 1000))
            for turn, _, _ in diar.itertracks(yield_label=True)
        ]
        yield {"phase":"diarization","status":"end","count":len(segments)}

    # 3) transcription en parallèle
    yield {"phase":"transcription","total":len(segments),"done":0}
    texts = [None] * len(segments)
    with ThreadPoolExecutor(max_workers=4) as ex:
        futures = {}
        for i, (start_ms, end_ms) in enumerate(segments):
            seg = audio_seg[start_ms:end_ms] if end_ms else audio_seg
            futures[ex.submit(_encode_and_transcribe, seg)] = i

        done = 0
        for fut in as_completed(futures):
            idx = futures[fut]
            texts[idx] = fut.result()
            done += 1
            yield {"phase":"transcription","done":done}

    # 4) reconstruction du transcript avec/sans locuteurs
    if pipeline and duration_ms <= _DIAR_THRESHOLD:
        # on reprend l’ordre des locuteurs
        speakers = [label for _,_,label in diar.itertracks(yield_label=True)]
        transcript = "\n".join(f"[{speakers[i]}] {texts[i]}" for i in range(len(texts)))
    else:
        transcript = texts[0]

    # 5) résumé
    yield {"phase":"summary","status":"start"}
    summary_resp = openai.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role":"system","content":(
            "Tu es un assistant expert en rédaction de rapports de réunion. "
            "Ton rôle est de transformer une transcription brute, étiquetée par locuteur, "
            "en un document professionnel, clair et structuré."
        )},
            {"role":"user","content":("Voici la transcription complète de la réunion, avec les étiquettes de locuteurs :\n\n"
            f"{transcript}\n\n"
            "Tu es un expert en analyse de réunion. À partir de la transcription suivante, rédige un rapport structuré et très détaillé. "
"Consignes :"
"1. Identifie clairement chaque interlocuteur (Prénom ou identifiant, s’il est précisé)."
"2. Résume et détaille précisément ce que chaque interlocuteur a dit, point par point, en respectant l’ordre chronologique."
"3. Distingue les interventions, les idées principales, les arguments, les questions, les réponses, les décisions prises, les désaccords éventuels, et les actions à suivre.comme ceci a),b) etc.."
"4. N’omets aucun sujet abordé, même brièvement., les Titres doit etre en gras "
"5. Utilise des titres, sous-titres et puces pour une lecture claire et professionnelle. n'utilise pas de ** ou hastage pour le titre juste 1), 2) etc."
"6. Termine le rapport par une section “Synthèse & prochaines étapes” regroupant :"
   "- Les points clés abordés. comme ceci a),b) etc.."
  " - Les décisions prises. comme ceci a),b) etc.."
   "- Les tâches/action points identifiés (avec responsables si mentionnés). comme ceci a),b) etc.."
            
            
            )}
        ],
        max_tokens=1500,
    )
    summary = summary_resp.choices[0].message.content
    yield {"phase":"summary","status":"end"}

    # 6) génération du .docx
    yield {"phase":"docx","status":"start"}
    docx_path = os.path.join("recordings", f"{os.path.basename(audio_file)}.report.docx")
    doc = Document()
    doc.add_heading("Transcription détaillée", level=1)
    doc.add_paragraph(transcript)
    doc.add_page_break()
    doc.add_heading("Synthèse de la réunion", level=1)
    doc.add_paragraph(summary)
    doc.save(docx_path)
    yield {"phase":"docx","status":"end","path":docx_path}

    # cleanup
    os.remove(tmp_wav)

    # return final values
    return transcript, summary, docx_path

# ——————————————————————————————
# Fonctions de secours (non utilisées dans SSE)
# ——————————————————————————————
def summarize_text(text: str) -> str:
    """
    Résumé rapide sans progression.
    """
    resp = openai.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role":"system","content":("Tu es un expert en analyse de réunion. À partir de la transcription suivante, rédige un rapport structuré et très détaillé."
"Consignes :"
"1. Identifie clairement chaque interlocuteur (Prénom ou identifiant, s’il est précisé)."
"2. Résume et détaille précisément ce que chaque interlocuteur a dit, point par point, en respectant l’ordre chronologique."
"3. Distingue les interventions, les idées principales, les arguments, les questions, les réponses, les décisions prises, les désaccords éventuels, et les actions à suivre."
"4. N’omets aucun sujet abordé, même brièvement."
"5. Utilise des titres, sous-titres et puces pour une lecture claire et professionnelle."
"6. Termine le rapport par une section “Synthèse & prochaines étapes” regroupant :"
   "- Les points clés abordés."
  " - Les décisions prises."
   "- Les tâches/action points identifiés (avec responsables si mentionnés).")},
            {"role":"user","content":text}
        ],
        max_tokens=1500,
    )
    return resp.choices[0].message.content

def generate_word(transcription: str,
                  summary: str,
                  output_doc: str = "recordings/meeting_report.docx") -> str:
    """
    Génère un .docx sans progression.
    """
    doc = Document()
    doc.add_heading("Transcription détaillée", level=1)
    doc.add_paragraph(transcription)
    doc.add_page_break()
    doc.add_heading("Synthèse de la réunion", level=1)
    doc.add_paragraph(summary)
    doc.save(output_doc)
    return output_doc
